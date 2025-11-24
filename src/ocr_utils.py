"""
Ce fichier contient des fonctions utilitaires pour le traitement d'images et
la reconnaissance de caractères.

Auteurs :
Pierrick BERTHE
mail : pierrick.berthe@gmx.fr
Août 2025
"""
# Imports standard
import datetime
import sys
import os
import time
from functools import wraps
import json
import re
import shutil
import subprocess

# Import third-party libraries
import language_tool_python
import cv2
from PIL import Image
import numpy as np
import easyocr
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import sqlite3

# pylint: disable=no-member

class Logger(object):
    """Logger class to redirect print statements to a file.
    """
    def __init__(self, log_path):
        self.terminal = sys.stdout
        self.log = open(log_path, "a", encoding="utf-8")
    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)
    def flush(self):
        pass


def get_git_version():
    """
    Get the current git version.
    """
    try:
        # Get the current git version
        version = subprocess.check_output(
            ["git", "describe", "--tags", "--always"],
            stderr=subprocess.STDOUT
        ).decode().strip()
        return version
    except Exception:
        return "version inconnue"


def format_git_version(version_str):
    """
    Format the git version string.
    """
    # split the version string
    parts = version_str.split('-')

    # Check the number of parts
    if len(parts) == 1:
        return f"Version : {parts[0]}"
    elif len(parts) == 3:
        tag, commits, commit_hash = parts
        return (
            f"Version : {tag} ({commits} commits après le tag,"
            f" commit {commit_hash})"
        )
    else:
        return f"Version : {version_str}"


def check_and_create_directories(*dirs):
    """
    check if directories exist, if not, create them.
    """
    for d in dirs:
        if not os.path.exists(d):
            os.makedirs(d, exist_ok=True)


def get_user_inputs(data_dir):
    """
    Ask the user for inputs such as client acronym, date of measurement,
    and folder to ignore. Returns these inputs.
    """
    # create acronym for the client name
    client_acronym = input("Entrez les initiales du nom du client : ")

    # Define the date of the measurement (format: JJ/MM/AAAA and before today)
    while True:
        try:
            date_mesure = input(
                "Entrez la date de la mesure (format JJ/MM/AAAA) : "
            )
            date_obj = datetime.datetime.strptime(date_mesure, "%d/%m/%Y")
            if date_obj > datetime.datetime.now():
                print("La date ne peut pas être dans le futur.")
                continue
            break
        except ValueError:
            print(
                "Format incorrect ou date impossible. "
                "Veuillez entrer la date au format JJ/MM/AAAA (ex. 30/06/2025)."
            )

    # Specify the folder to ignore
    folder_ignored = input("Entrez le nom exact du dossier à ignorer : ")
    folder_ignored_dir = os.path.join(data_dir, folder_ignored)

    return client_acronym, date_mesure, folder_ignored, folder_ignored_dir


def get_files_by_subdir(folder_path, folder_source_name="source"):
    """
    Get all files in the subdirectories of the given folder.
    """
    # Initialize a dictionary to hold the files by subdirectory
    files_by_subdir = {}

    # FOR each subdirectory in the folder, get the filenames
    for root, dirs, filenames in os.walk(folder_path):
        if root == folder_path:
            continue
        files_by_subdir[root.split(folder_source_name + "\\")[-1]] = filenames
    return files_by_subdir


def preprocess_black_text(image_path, output_path):
    """
    Preprocess the image to make all non-black pixels white, highlighting
    black text. Uses PIL for image loading.
    """
    try:
        # Load the image with PIL and convert to RGB then to numpy array
        with Image.open(image_path) as img:
            img = img.convert("RGB")
            image = np.array(img)
    except Exception as e:
        print(f"Erreur : Impossible de charger l'image {image_path} : {e}")
        return None

    # Convert the image to HSV color space
    hsv = cv2.cvtColor(image, cv2.COLOR_RGB2HSV)

    # Define the lower and upper bounds for black color in HSV
    lower_black = np.array([0, 0, 0])
    upper_black = np.array([0, 0, 0]) 

    # Create a mask for black pixels
    mask = cv2.inRange(hsv, lower_black, upper_black)

    # Apply the mask to the image
    result_image = cv2.bitwise_and(image, image, mask=mask)

    # Replace non-black pixels with white
    result_image[np.where(mask == 0)] = [255, 255, 255]

    # Save the preprocessed image
    cv2.imwrite(output_path, result_image)

    return result_image


def resize_image(input_path, output_path, scale_percent=10):
    """
    Resize an image to a given percentage of its original size.
    """
    try:
        # Load the image with PIL and convert to RGB then to numpy array
        with Image.open(input_path) as img:
            img = img.convert("RGB")
            image = np.array(img)
    except Exception as e:
        print(f"Erreur : Impossible de charger l'image {input_path} : {e}")
        return None

    # Calculate the new dimensions
    width = int(image.shape[1] * scale_percent / 100)
    height = int(image.shape[0] * scale_percent / 100)
    dim = (width, height)

    # Resize the image
    resized_image = cv2.resize(image, dim, interpolation=cv2.INTER_AREA)

    # Save the resized image
    cv2.imwrite(output_path, resized_image)

    return resized_image


def measure_time(func):
    """
    Decorator to measure the execution time of a function.
    """
    @wraps(func)
    def wrapper(*args, **kwargs):
        start = time.time()
        result = func(*args, **kwargs)
        duration = time.time() - start
        return result, duration
    return wrapper


@measure_time
def extract_text_easyocr(
    image_path,
    batch_size,
    decoder,
    adjust_contrast,
    worker,
    gpu_state
):
    """
    Extract text from image using EasyOCR
    """
    # List of character
    french_characters = (
        "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
        "àâäéèêëîïôöùûüçÀÂÄÉÈÊËÎÏÔÖÙÛÜÇ' -"
    )

    # Execute OCR
    reader = easyocr.Reader(['fr'], gpu=gpu_state, verbose=False)
    text = reader.readtext(
        image_path,
        detail=0,
        batch_size=batch_size,
        decoder=decoder,
        adjust_contrast=adjust_contrast,
        workers=worker,
        allowlist=french_characters
    )
    return "\n".join(text)


@measure_time
def correct_text_french(text):
    """
    Correct the text in French using LanguageTool.
    """
    tool = language_tool_python.LanguageTool('fr')
    corrected = tool.correct(text)
    return corrected


def export_text_to_json(
        text_extracted,
        output_dir,
        output_file="text_extracted.json"
    ):
    """
    Export the text_extracted dictionary to a JSON file in the specified
    directory.
    """
    # Create the output path
    file_path = os.path.join(output_dir, output_file)

    try:
        # Save the dictionary to a JSON file
        with open(file_path, "w", encoding="utf-8") as json_file:
            json.dump(text_extracted, json_file, ensure_ascii=False, indent=4)
        print(f"Dictionnaire exporté dans le fichier :\n {file_path}\n")

    # Error handling
    except Exception as e:
        print(f"Erreur lors de l'exportation : {e}")


def clean_text(text):
    '''
    Clean the text by removing unwanted characters and formatting.
    '''

    # If text is a tuple or not a string, convert it to a string
    if isinstance(text, tuple):
        text = " ".join(str(t) for t in text)
    elif not isinstance(text, str):
        text = str(text)

    # Replace line break (\n) and carriage return (-r) with space
    text = re.sub(r"[\n\r]", " ", text)

    # Keep only letters & digits (\w), spaces (\s), and dashes (-)
    text = re.sub(r"[^\w\s'-]", "", text)

    # Replace multiple spaces (\s+) with a single space
    text = re.sub(r"\s+", " ", text)

    # Remove trailing spaces and capitale letter
    return text.strip().lower()


def extract_key_info(text, pattern=r"[a-zA-Z]\d+"):
    """
    Extract key information from the text.
    """
    # Search for the pattern in the text
    matches = list(re.finditer(pattern, text))

    if matches:

        # Get the start and end indices of the match
        first_start = matches[0].start()
        client_name = text[:first_start].strip()

        # Extract all chimney names and remarks
        chimney_names = []
        remarks = text[matches[-1].end():].strip()

        for match in matches:
            chimney_names.append(match.group())

        return {
            "client_name": client_name,
            "chimney_name": chimney_names,
            "remarks": remarks
        }
    else:

        # Log a warning if no match is found
        print(f"\nWarning: No chimney found in the text: {text}")

        # If no match is found, return None for chimney name and remarks
        return {
            "client_name": "",
            "chimney_name": "",
            "remarks": ""
        }


def generate_filename(
        text,
        pattern,
        client_acronym,
        filenames,
        extension="jpg"
    ):
    """
    Generate filenames based on the extracted text.
    If a filename already exists in 'filenames', increment the counter like Windows.
    """
    # create a empty list to hold the new filenames
    new_filenames = []

    # Clean the text and extract key information
    cleaned_text = clean_text(text)
    key_info = extract_key_info(cleaned_text, pattern)

    # If no chimney name is found, return "Same_as_original"
    if not key_info["chimney_name"]:
        new_filenames.append("Same_as_original")
        return new_filenames, key_info

    # FOR each chimney name, generate a unique filename
    for chimney_name in key_info["chimney_name"]:

        # create the base name and the regex pattern to match existing filenames
        base_name = f"{client_acronym}_{chimney_name}"
        pattern = re.compile(
            rf"^{re.escape(base_name)}(?: \((\d+)\))?\.{re.escape(extension)}$"
        )
        max_counter = 0

        # Search for existing filenames and find the maximum counter
        for fname in filenames:
            match = pattern.match(fname)
            if match:
                # If a counter is found, update max_counter
                if match.group(1):
                    num = int(match.group(1))
                    if num > max_counter:
                        max_counter = num
                # if no counter is found, consider counter=0
                else:
                    if max_counter == 0:
                        max_counter = 0

        # Generate the next name
        if max_counter == 0 and f"{base_name}.{extension}" not in filenames:
            new_filename = f"{base_name}.{extension}"
        else:
            new_filename = f"{base_name} ({max_counter + 1}).{extension}"

        # Append the new filename to the list
        new_filenames.append(new_filename)

    return new_filenames, key_info


def copy_files_with_mapping(
        text_extracted,
        pattern,
        input_dir,
        output_folder_dir,
        output_json_dir,
        client_acronym,
        key_info_file="key_info.json",
        mapping_file="file_mapping.json"
    ):
    """
    Copy files from the input directory to the output directory with new
    names based on extracted text.
    """
    # Initialise dictionaries for the mapping and key information
    mapping, key_info_dict = {}, {}

    # Create the output directory if it doesn't exist
    for subdir, files in text_extracted.items():
        subdir = subdir.split("\\")[-1]
        output_subdir = os.path.join(output_folder_dir, subdir)
        os.makedirs(output_subdir, exist_ok=True)
        mapping[subdir], key_info_dict[subdir] = {}, {}
        filenames = []
        print(f"\nSUBDIRECTORY: {subdir}")
        print("-" * 75)

        # FOR each file, generate new filenames
        for file, text in files.items():
            new_filenames, key_info = generate_filename(
                text,
                pattern,
                client_acronym,
                filenames=filenames
            )

            # append the new filenames to the list
            filenames.extend(new_filenames)

            # Store the key information in the dictionary
            key_info_dict[subdir][file] = key_info

            # Paths for the old file
            old_path = os.path.join(input_dir, subdir, file)

            # Copy the file to the new location with each generated name
            copied_files = []
            for new_filename in new_filenames:
                # Use the original filename if "Same_as_original" is returned
                if new_filename == "Same_as_original":
                    new_filename = file

                # Generate a new path for the copied file
                new_path = os.path.join(output_subdir, new_filename)
                shutil.copy(old_path, new_path)
                copied_files.append(new_filename)

                # Print each copied file on a separate line
                print(f"Copied: {file} -> {new_filename}")

            # Add the mapping to the dictionary
            mapping[subdir][file] = copied_files

    # Save key information and mapping to JSON files
    print("\nExports des fichiers JSON : ")
    export_text_to_json(key_info_dict, output_json_dir, key_info_file)
    export_text_to_json(mapping, output_json_dir, mapping_file)

    return key_info_file, mapping_file


def import_json_to_text(input_dir, input_file="text_extracted.json"):
    """
    Import the JSON file in the specified directory.
    """
    # Create the input path
    file_path = os.path.join(input_dir, input_file)

    try:
        # Load the JSON file to dicctionnary
        with open(file_path, "r", encoding="utf-8") as json_file:
            data = json.load(json_file)
        print(f"Dictionnaire importé depuis le fichier :\n{file_path}\n")
        return data

    # Error handling
    except FileNotFoundError:
        print(f"Erreur : Le fichier '{file_path}' n'existe pas.")
    except json.JSONDecodeError:
        print(
            f"Erreur : Le fichier '{file_path}' n'est pas "
            "un fichier JSON valide."
        )
    except Exception as e:
        print(f"Erreur inattendue lors de l'importation du fichier JSON : {e}")


def sort_key(chimney_name, pattern):
    """
    Extract all alphabetical and numerical parts of the chimney name
    for sorting.
    """
    # Find all matches of the pattern in the chimney name
    matches = re.findall(pattern, chimney_name)
    if matches:
        # If match is a tuple (from groups), flatten if needed
        if isinstance(matches[0], tuple):
            sorted_parts = [(m[0], int(m[1])) for m in matches]
        else:
            # fallback: treat as string, not tuple
            sorted_parts = [(matches[0], 0)]
        return sorted_parts
    return [(chimney_name, 0)]


def group_by_chimney_name(data, pattern=r"[a-zA-Z]\d+"):
    """
    Group the data by chimney name and return a dictionary sorted by chimney
    name in alphabetical order and numerical value. 
    """
    # Initialize an empty dictionary to hold the grouped data
    grouped_data = {}
    counter = 1

    # Iterate through the data and group by chimney name
    for subdir, files in data.items():
        for file, info in files.items():
            chimney_names = info.get("chimney_name", "")

            # Ensure chimney_names is a list
            if not isinstance(chimney_names, list):
                chimney_names = [chimney_names]

            # Handle the case where chimney_names is empty
            if not chimney_names or chimney_names == [""]:
                chimney_names = [f"No_chimney_{counter}"]
                counter += 1

            # Iterate through each chimney name and group the data
            for chimney_name in chimney_names:
                if chimney_name not in grouped_data:
                    grouped_data[chimney_name] = []
                grouped_data[chimney_name].append({
                    "subdir" : subdir,
                    "file": file,
                    "client_name": info["client_name"],
                    "remarks": info["remarks"]
                })

    # Sort the dictionary by the custom sort key
    grouped_data = dict(sorted(
        grouped_data.items(),
        key=lambda item: sort_key(item[0], pattern)
    ))

    return grouped_data


def save_to_json(data, output_dir, filename):
    """
    Save data to a JSON file in the specified output directory.
    """
    # create the file path
    file_path = os.path.join(output_dir, filename)

    # save the data to JSON file
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print(f"\nDonnées sauvegardées dans le fichier JSON :\n {file_path}")
        return file_path

    # Error handling
    except FileNotFoundError:
        print(f"Erreur : Le répertoire de sortie '{output_dir}' n'existe pas.")
    except PermissionError:
        print(
            f"Erreur : Permission refusée pour écrire "
            f"dans le répertoire '{output_dir}'."
        )
    except Exception as e:
        print(f"Erreur inattendue lors de la sauvegarde du fichier JSON : {e}")
    finally:
        # Close the file if it was opened
        if 'f' in locals():
            f.close()


def get_client_name_counts(data):
    """
    Retrieves the count of occurrences for each client name using Pandas,
    ignoring empty client names.
    """
    # Flatten the data into a list of client names, ignore empty strings
    client_names = [
        info["client_name"]
        for subdir, files in data.items()
        for file, info in files.items()
        if info["client_name"] != ""
    ]

    # Create a Pandas Series and count occurrences
    client_name_series = pd.Series(client_names)
    client_name_counts = client_name_series.value_counts()

    return client_name_counts.to_dict()


def compress_image(
        image_path,
        temp_dir,
        max_width=800,
        max_height=800,
        quality=50
    ):
    """
    Reduce picture size using dimension et quality reduction
    """
    # Load picture with PIL
    try:
        with Image.open(image_path) as img:
            img = img.convert("RGB")
            image = np.array(img)
    except Exception as e:
        print(f"Erreur : Impossible de charger l'image {image_path} : {e}")
        return None

    # Obtain picture dimension
    height, width = image.shape[:2]

    # Calculate dimension factor
    scale = min(max_width / width, max_height / height, 1.0)
    new_width = int(width * scale)
    new_height = int(height * scale)

    # Resize picture
    resized_image = cv2.resize(
        image,
        (new_width, new_height),
        interpolation=cv2.INTER_AREA
    )

    # Convert BGR -> RGB before saving
    resized_image_rgb = cv2.cvtColor(resized_image, cv2.COLOR_RGB2BGR)

    # Temporary file path fo compressed picture
    compressed_image_path = os.path.join(temp_dir, "temp_compressed_image.jpg")

    # Save compressed picture with reducted quality
    cv2.imwrite(
        compressed_image_path,
        resized_image_rgb,
        [cv2.IMWRITE_JPEG_QUALITY, quality]
    )

    return compressed_image_path


def add_page_number_field(paragraph):
    """
    Add a page number field to a Word document paragraph.
    """
    run = paragraph.add_run("Page ")

    # Add PAGE field (current page number)
    fldChar_begin = docx.oxml.OxmlElement('w:fldChar')
    fldChar_begin.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
    instrText = docx.oxml.OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar_end = docx.oxml.OxmlElement('w:fldChar')
    fldChar_end.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)
    run.add_text(" / ")

    # Add NUMPAGES field (total number of pages)
    fldChar_begin2 = docx.oxml.OxmlElement('w:fldChar')
    fldChar_begin2.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
    instrText2 = docx.oxml.OxmlElement('w:instrText')
    instrText2.text = "NUMPAGES"
    fldChar_end2 = docx.oxml.OxmlElement('w:fldChar')
    fldChar_end2.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin2)
    run._r.append(instrText2)
    run._r.append(fldChar_end2)


def add_picture_to_paragraph(paragraph, image_path, width=Inches(1.5)):
    """
    Add a picture to a Word document paragraph, or text if the image is
    unavailable.
    """
    run = paragraph.add_run()
    try:
        if image_path is not None:
            run.add_picture(image_path, width=width)
        else:
            paragraph.add_run("Image non disponible")
    except Exception as e:
        print(
            "Erreur lors de l'ajout de l'image dans le "
            f"rapport Word : {e}"
        )
        paragraph.add_run("Image non disponible")


def generate_word_report(
        data_per_chimney,
        input_dir,
        output_dir,
        temp_dir,
        client_name,
        files_by_subdir,
        date_mesure,
        logo_path,
        output_file_name="Annexes photo 3CEP.docx"
    ):
    """
    Generate a Word report containing compressed images, extracted information,
    files_by_subdir keys, and a logo in the header.
    """
    # Create a Word document
    document = Document()

    # Add a header with the logo
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add the logo to the header
    if os.path.exists(logo_path):
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.0))

    # Add the right-aligned header text
    right_header_paragraph = header.add_paragraph()
    right_header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    right_header_paragraph.add_run("Annexe inspection R'PUR CC").bold = True

    # Add the main centered title
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(
        "Rapport photos inspection 3CEP\n"
        f"{client_name.capitalize()}\n{date_mesure}"
    )
    run.bold = True
    run.font.size = Pt(16)

    # Add footer with company information and page numbers
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.add_run(
        "SARL R'PUR Conduits Collectifs \nLe présent rapport rend comptes des"
        " éléments vus, visitables et déclarés par l'exploitant.\n")
    add_page_number_field(footer_paragraph)

    # ensure temp_dir exists
    os.makedirs(temp_dir, exist_ok=True)

    # Flag to control page breaks for first chimney
    first_chimney = True

    # Loop through chimneys and their information
    for chimney, subdirs in data_per_chimney.items():

        # Skip if the chimney name is "No_chimney"
        if "No_chimney" in chimney:
            continue

        # start a new page for each chimney except the first
        if not first_chimney:
            document.add_page_break()

        # Add the chimney name
        title_1 = document.add_paragraph()
        title_1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = title_1.add_run(f"Conduit: {chimney}")
        run.bold = True
        run.underline = True
        run.font.size = Pt(14)        

        # Create a table with 2 columns: Location, Photo
        n_rows = len(files_by_subdir.keys())
        table = document.add_table(rows=n_rows, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(2.5)

        # Fill the first column with the names from files_by_subdir
        keys_list = list(files_by_subdir.keys())
        for i, key in enumerate(keys_list):
            split_key = key.split("\\")[-1]
            table.cell(i, 0).text = split_key

        # FOR each key in files_by_subdir (folders)
        for i, key in enumerate(keys_list):

            # Check if the key contains backslashes and split accordingly
            split_key = key.split("\\")[-1] if "\\" in key else key
            cell = table.cell(i, 1)
            images_found = False

            # FOR each list of subdirs for the current chimney
            for entry in subdirs:
                folder_name = entry["subdir"]

                # Check if the split key exists in subdirs
                if split_key in folder_name:
                    file = entry["file"]

                    # Photo column
                    image_path = os.path.join(input_dir, key, file)
                    if os.path.exists(image_path):
                        compressed_image_path = compress_image(
                            image_path, temp_dir
                        )
                        paragraph = cell.add_paragraph()
                        add_picture_to_paragraph(
                            paragraph, compressed_image_path
                        )
                        paragraph.paragraph_format.space_before = Pt(4)
                        paragraph.paragraph_format.space_after = Pt(4)
                        images_found = True

                if not images_found:
                    cell.text = ""

        # Update the flag after processing the first chimney
        first_chimney = False

    # Build the full path for the output file
    output_file_path = os.path.join(output_dir, output_file_name)

    # Save the Word document
    document.save(output_file_path)
    print(f"\nRapport généré:\n {output_file_path}")

    return output_file_path


def create_database_and_tables(bdd_dir, db_name="bdd_airpur.db"):
    """
    Create a SQLite database and its tables. 
    If the database already exists, it is not recreated.
    """
    # Build the database file path
    bdd_path = os.path.join(bdd_dir, db_name)

    # If the database file does not exist, create it and the tables
    if not os.path.exists(bdd_path):

        # Connect to the SQLite database
        conn = sqlite3.connect(bdd_path, timeout=10)
        cur = conn.cursor()

        # Create the 'clients' table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS clients (
            client_id TEXT PRIMARY KEY,
            nom TEXT
        )
        """)

        # Create the 'cheminees' table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS cheminees (
            client_id TEXT,
            cheminee_id TEXT,
            localisation TEXT,
            remarques TEXT,
            PRIMARY KEY (cheminee_id, client_id, localisation),
            FOREIGN KEY (client_id) REFERENCES clients(client_id)
        )
        """)

        # Create the 'mesures' table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS mesures (
            client_id TEXT,
            cheminee_id INTEGER,
            mesure_id INTEGER,
            date_mesure DATETIME,
            PRIMARY KEY (mesure_id, client_id, cheminee_id),
            FOREIGN KEY (cheminee_id) REFERENCES cheminees(cheminee_id),
            FOREIGN KEY (client_id) REFERENCES clients(client_id)
        )
        """)

        # Commit changes and close the connection
        conn.commit()
        conn.close()
        print(f"\nBase de données (BDD) créée :\n {bdd_path}")
    else:
        print(f"\nBase de données (BDD) déjà existante :\n {bdd_path}")
    print("\n" + "." * 75)

    return bdd_path


def insert_client_into_db(bdd_path, client_acronym, client_name):
    """
    Insert a client into the SQLite database.
    If the client already exists, it will not be inserted again.
    Displays the total number of clients and the number of clients added.
    """
    print("\nTABLE clients : ")
    client_added = False
    try:
        # Connect to the database
        conn = sqlite3.connect(bdd_path, timeout=10)
        cur = conn.cursor()

        # Insert client data into the table
        try:
            cur.execute(
                "INSERT INTO clients (client_id, nom) VALUES (?, ?)",
                (client_acronym, client_name)
            )
            conn.commit()
            print(
                f"Client '{client_name}', acronym '{client_acronym}' : "
                f"ajouté à la BDD.\n"
            )
            client_added = True

        # If the client already exists, print a message
        except sqlite3.IntegrityError:
            print(
                f"Client '{client_name}', acronym '{client_acronym}' : "
                f"existant déjà dans la BDD.\n"
            )

        # Print error if insertion fails
        except Exception as e:
            print(f"Erreur lors de l'insertion du client : {e}\n")

        # Count the total number of clients in the database
        cur.execute("SELECT COUNT(*) FROM clients")
        total_clients = cur.fetchone()[0]
        print(f"Nombre de client ajouté dans BDD : {1 if client_added else 0}")
        print(f"Nombre total de clients dans la BDD : {total_clients}")
        print("\n" + "." * 75)

    # Error handling for connection issues
    except Exception as e:
        print(f"Erreur de connexion ou d'exécution SQL : {e}")

    # Ensure the connection is closed
    finally:
        try:
            conn.close()
        except Exception as e:
            print(f"Erreur lors de la fermeture de la connexion : {e}")


def insert_cheminee_into_db(bdd_path, client_acronym, data_per_chimney):
    """
    Insertion of chimneys into the SQLite database.
    If the chimney already exists, it will not be inserted again.
    Displays the total number of chimneys and the number of chimneys added.
    """
    print("\nTABLE cheminées : ")
    cheminees_added = 0
    try:
        # Connect to the database
        conn = sqlite3.connect(bdd_path, timeout=10)
        cur = conn.cursor()

        # FOR each chimney in the data_per_chimney dictionary
        for cheminee_id, entries in data_per_chimney.items():

            # Skip if the chimney name is "No_chimney"
            if "No_chimney" in cheminee_id:
                continue

            # Insert each subdirectory as a chimney using localisation
            for entry in entries:
                subdir = entry["subdir"]
                remark = entry["remarks"]

                # Insert the chimney data into the table
                try:
                    cur.execute(
                        "INSERT INTO cheminees "
                        "(client_id, cheminee_id, localisation, remarques) "
                        "VALUES (?, ?, ?, ?)",
                        (client_acronym, cheminee_id, subdir, remark)
                    )
                    cheminees_added += 1
                    print(
                        f"Client: '{client_acronym}', "
                        f"Cheminée: '{cheminee_id}', "
                        f"Localisation: '{subdir}' : "
                        "ajoutée à la BDD."
                    )

                # Exception handling for duplicate entries
                except sqlite3.IntegrityError:
                    print(
                        f"Client: '{client_acronym}', "
                        f"Cheminée: '{cheminee_id}', "
                        f"Localisation: '{subdir}' : "
                        "existe déjà dans la BDD."
                    )
                except Exception as e:
                    print(
                        "Erreur lors de l'insertion de la cheminée "
                        f"{cheminee_id} : {e}"
                    )
        conn.commit()

        # Count the total number of chimneys in the database
        cur.execute("SELECT COUNT(*) FROM cheminees")
        total_cheminees = cur.fetchone()[0]
        print(
            "\nNombre de combinaison client/cheminée/localisation "
            f"ajoutée(s) dans la BDD : {cheminees_added}"
        )
        print(
            "Nombre total de combinaison client/cheminée/localisation "
            f"dans la BDD : {total_cheminees}"
        )
        print("\n" + "." * 75)

    # Error handling for connection issues
    except Exception as e:
        print(f"Erreur de connexion ou d'exécution SQL : {e}")

    # Ensure the connection is closed
    finally:
        try:
            conn.close()
        except Exception as e:
            print(f"Erreur lors de la fermeture de la connexion : {e}")


def insert_mesure_into_db(
        bdd_path,
        client_acronym,
        data_per_chimney,
        date_mesure
    ):
    """
    Insertion of measurements into the SQLite database.
    If the measurement already exists, it will not be inserted again.
    Checks the last mesure_id and date_mesure for each client and chimney.
    Displays the total number of measurements and the number of measurements
    added.
    """
    print("\nTABLE mesures : ")
    mesures_added = 0
    try:
        # Connect to the database
        conn = sqlite3.connect(bdd_path, timeout=3)
        cur = conn.cursor()

        # FOR each chimney in the data_per_chimney dictionary
        for cheminee_id, entries in data_per_chimney.items():

            # Skip if the chimney name is "No_chimney"
            if "No_chimney" in cheminee_id:
                continue

            # FOR each subdirectory in the chimney
            for entry in entries:

                # Format the measurement date to SQL format
                date_mesure_str = date_mesure
                try:
                    date_obj = datetime.datetime.strptime(
                        date_mesure_str, "%d/%m/%Y"
                    )
                    date_mesure_sql = date_obj.strftime("%Y-%m-%d")
                except Exception:
                    date_mesure_sql = date_mesure_str

                # Find the last mesure_id and date for this client and chimney
                cur.execute("""
                    SELECT MAX(mesure_id), MAX(date_mesure) FROM mesures
                    WHERE client_id = ? AND cheminee_id = ?""",
                    (client_acronym, cheminee_id)
                )
                row = cur.fetchone()
                last_mesure_id, last_date_in_db = row if row else (None, None)

                # If no previous measure, start at 1
                if last_mesure_id is None:
                    mesure_id = 1
                else:
                    # If the current date is newer, increment mesure_id
                    if (
                        last_date_in_db is not None
                        and date_mesure_sql > last_date_in_db
                    ):
                        mesure_id = last_mesure_id + 1
                    else:
                        mesure_id = last_mesure_id

                # Check if the measure already exists
                cur.execute("""
                    SELECT COUNT(*) FROM mesures
                    WHERE client_id = ? AND cheminee_id = ? 
                    AND mesure_id = ?""",
                    (client_acronym, cheminee_id, mesure_id)
                )
                exists = cur.fetchone()[0]

                # If the measure already exists, skip to the next
                if exists:
                    print(
                        f"Client: '{client_acronym}', "
                        f"Cheminée: '{cheminee_id}', "
                        f"date de mesure: '{date_mesure_sql}', "
                        f"mesure_id: {mesure_id} : "
                        "déjà existante dans la BDD."
                    )
                    continue

                # Insert the measure into the table
                try:
                    cur.execute(
                        "INSERT INTO mesures "
                        "(client_id, cheminee_id, mesure_id, date_mesure) "
                        "VALUES (?, ?, ?, ?)",
                        (
                            client_acronym,
                            cheminee_id,
                            mesure_id,
                            date_mesure_sql
                        )
                    )
                    mesures_added += 1
                    print(
                        f"Client: '{client_acronym}', "
                        f"Cheminée: '{cheminee_id}', "
                        f"date de mesure: '{date_mesure_sql}', "
                        f"mesure_id: {mesure_id} : ajoutée à la BDD."
                    )

                # Exception handling for duplicate entries
                except Exception as e:
                    print(
                        "Erreur lors de l'insertion de la mesure pour "
                        f"{cheminee_id} : {e}"
                    )
        conn.commit()

        # Count the total number of measures in the database
        cur.execute("SELECT COUNT(*) FROM mesures")
        total_mesures = cur.fetchone()[0]
        print(
            "\nNombre de combinaison client/cheminée/date ajoutée(s) "
            f"dans la BDD : {mesures_added}"
        )
        print(
            "Nombre total de combinaison client/cheminée/date dans la BDD : "
            f"{total_mesures}"
        )
        print("\n" + "." * 75)

    # Error handling for connection issues
    except Exception as e:
        print(f"Erreur de connexion ou d'exécution SQL : {e}")

    # Ensure the connection is closed
    finally:
        try:
            conn.close()
        except Exception as e:
            print(f"Erreur lors de la fermeture de la connexion : {e}")


def print_step(step_num, message):
    """
    Print a formatted step message for process reporting.
    """
    print("\n\n" + "*" * 75)
    print(f"==> STEP {step_num} : {message} <==")
    print("*" * 75)


def calculate_duration(start_time):
    """
    This function calculates and displays the execution duration of scripts.
    It prints the duration in minutes and seconds.
    """
    minutes, seconds = divmod(time.time() - start_time, 60)
    print(f"\nDurée execution script : {int(minutes)} min {int(seconds)} sec")


def backup_database(
    db_path,
    backup_dir,
    max_backups=10
):
    """
    Create a timestamped backup of the SQLite database and keep only the
    latest N backups.
    """
    # Ensure the backup directory exists
    os.makedirs(backup_dir, exist_ok=True)

    # Create a timestamp for the backup filename
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")

    # Get the database name without extension
    db_name = os.path.basename(db_path).split(".")[0]

    # Build the backup filename
    backup_file = os.path.join(backup_dir, f"{db_name}_backup_{timestamp}.db")

    # Copy the database file to the backup location
    shutil.copy(db_path, backup_file)
    print(f"\n[INFO] Backup de la base de données créée :\n{backup_file}\n")

    # Sort and keep only the N most recent backups
    backups = sorted(
        [
            f for f in os.listdir(backup_dir)
            if f.startswith(f"{db_name}_backup_")
        ],
        key=lambda x: os.path.getctime(os.path.join(backup_dir, x))
    )
    if len(backups) > max_backups:
        old_backups = backups[:-max_backups]
        for old in old_backups:
            os.remove(os.path.join(backup_dir, old))
            print(f"[INFO] Old backup deleted:\n{old}\n")


def backup_word_report(
    word_path,
    backup_dir,
    client_acronym
):
    """
    Create a timestamped backup of the Word report.
    """
    # Ensure the backup directory exists
    os.makedirs(backup_dir, exist_ok=True)

    # Create a timestamp for the backup filename
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")

    # Build the backup filename
    backup_file = os.path.join(
        backup_dir,
        f"word_report_{client_acronym}_{timestamp}.docx"
    )
    # Copy the Word report to the backup location
    shutil.copy(word_path, backup_file)
    print(f"[INFO] Backup du rapport Word créé :\n{backup_file}\n")


def backup_json(
    json_path,
    backup_dir,
    client_acronym
):
    """
    Create a timestamped backup of the JSON files of client.
    """
    # Ensure the backup directory exists
    os.makedirs(backup_dir, exist_ok=True)

    # Create a timestamp for the backup filename
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")

    # Create a folder for the client acronym and timestamp
    folder_name = f"json_{client_acronym}_{timestamp}"
    client_backup_dir = os.path.join(backup_dir, folder_name)
    os.makedirs(client_backup_dir, exist_ok=True)

    # Copy all JSON files in the json_path to the client backup directory
    for file_name in os.listdir(json_path):
        if file_name.endswith(".json"):
            src_file = os.path.join(json_path, file_name)
            dst_file = os.path.join(client_backup_dir, file_name)
            shutil.copy(src_file, dst_file)
            print(f"[INFO] Backup du fichier JSON créé :\n{dst_file}\n")
    print(f"[INFO] Tous les fichiers JSON ont été sauvegardés dans :\n{client_backup_dir}\n")
